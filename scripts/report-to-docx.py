#!/usr/bin/env python3
"""
Convert the mid-year report Markdown draft into an IEEE single-column Word (.docx).

Handles the Markdown subset this draft uses: ATX headings, inline **bold** /
*italic* / `code`, pipe tables, images (with their figure captions), blockquotes,
horizontal rules, and bullet/numbered lists. Not a general Markdown engine —
scoped to docs/report/results-discussion-conclusion-draft.md.

Styling follows the IEEE single-column manuscript convention: Times New Roman,
10 pt justified body, bold numbered headings, centred title/author block, and
figures/tables sized to the text column.

Usage:  python3 scripts/report-to-docx.py [--uoa12] [input.md] [output.docx]

--uoa12 switches from the IEEE 10 pt/letter layout to the UoA final-report
template convention: A4 page, Times New Roman 12 pt body (tables and captions
at 10 pt). Inline <sub>/<sup> HTML tags are rendered as true sub/superscripts
in both modes.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
args = [a for a in sys.argv[1:] if a != "--uoa12"]
UOA12 = "--uoa12" in sys.argv[1:]
IN = (Path(args[0]) if len(args) > 0 else ROOT / "docs/report/results-discussion-conclusion-draft.md").resolve()
OUT = (Path(args[1]) if len(args) > 1 else ROOT / "docs/report/DementiaGuide_MidYear_Report.docx").resolve()

FONT = "Times New Roman"
BODY_PT = 12 if UOA12 else 10
SMALL_PT = max(BODY_PT - 2, 9)  # tables, captions
COL_WIDTH_IN = 6.27 if UOA12 else 6.5  # A4/letter width minus 1" margins each side

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
SUBSUP = re.compile(r"(<sub>.*?</sub>|<sup>.*?</sup>)")
IMG = re.compile(r"^!\[.*?\]\((.+?)\)\s*$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def set_font(run, size=None, bold=None, italic=None, mono=False, color=None):
    run.font.name = "Consolas" if mono else FONT
    # ensure east-asian/complex runs also use the font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), "Consolas" if mono else FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def emit_runs(paragraph, text, size, bold=None, italic=None, mono=False):
    """Emit runs for a plain-text span, honouring <sub>/<sup> HTML tags."""
    for part in SUBSUP.split(text):
        if not part:
            continue
        if part.startswith("<sub>"):
            r = paragraph.add_run(part[5:-6]); set_font(r, size, bold=bold, italic=italic, mono=mono)
            r.font.subscript = True
        elif part.startswith("<sup>"):
            r = paragraph.add_run(part[5:-6]); set_font(r, size, bold=bold, italic=italic, mono=mono)
            r.font.superscript = True
        else:
            r = paragraph.add_run(part); set_font(r, size, bold=bold, italic=italic, mono=mono)


def add_inline(paragraph, text, size=BODY_PT, base_bold=None, base_italic=None):
    """Render inline markdown (**bold**, *italic*, `code`) into a paragraph."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            emit_runs(paragraph, tok[2:-2], size, bold=True, italic=base_italic)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            emit_runs(paragraph, tok[1:-1], size - 0.5, italic=base_italic, mono=True)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            emit_runs(paragraph, tok[1:-1], size, bold=base_bold, italic=True)
        else:
            emit_runs(paragraph, tok, size, bold=base_bold, italic=base_italic)


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def main():
    lines = IN.read_text().splitlines()
    doc = Document()

    # Base document styling.
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    if UOA12:
        # Word's template default (~1.08 line spacing, 8 pt after) wastes ~15%
        # of the page; the UoA report convention is single-spaced.
        normal.paragraph_format.line_spacing = 1.0
        normal.paragraph_format.space_after = Pt(6)
    for sec in doc.sections:
        if UOA12:
            sec.page_width = Mm(210)
            sec.page_height = Mm(297)
        sec.top_margin = sec.bottom_margin = Inches(1)
        sec.left_margin = sec.right_margin = Inches(1)

    caption_re = re.compile(r"^\*\*(Figure|Table)\s")
    i = 0
    in_title_block = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped):
            in_title_block = False
            i += 1
            continue

        # heading
        m = HEADING.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                add_inline(p, text, size=18, base_bold=True)
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, text, size=BODY_PT + 2 if level == 2 else BODY_PT + 1, base_bold=True)
            i += 1
            continue

        # image  ->  centred picture (+ following caption paragraph)
        im = IMG.match(stripped)
        if im:
            path = (IN.parent / im.group(1)).resolve()
            if path.exists():
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(path), width=Inches(5.2 if UOA12 else COL_WIDTH_IN))
            i += 1
            continue

        # blockquote (possibly multi-line)
        if stripped.startswith(">"):
            buff = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buff.append(lines[i].strip()[1:].strip())
                i += 1
            text = " ".join(x for x in buff if x)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, text, size=BODY_PT - 1, base_italic=True)
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            header = split_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, cell_text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                add_inline(cell.paragraphs[0], cell_text, size=SMALL_PT, base_bold=True)
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncol):
                    cell = cells[c]
                    cell.paragraphs[0].text = ""
                    add_inline(cell.paragraphs[0], row[c] if c < len(row) else "", size=SMALL_PT)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # list item(s)
        if re.match(r"^(\-|\d+\.)\s+", stripped):
            while i < len(lines) and re.match(r"^(\-|\d+\.)\s+", lines[i].strip()):
                item = lines[i].strip()
                ordered = bool(re.match(r"^\d+\.", item))
                content = re.sub(r"^(\-|\d+\.)\s+", "", item)
                p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
                add_inline(p, content, size=BODY_PT)
                i += 1
            continue

        # plain paragraph (caption / author block get special alignment)
        p = doc.add_paragraph()
        if caption_re.match(stripped):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, stripped, size=SMALL_PT)
        elif in_title_block and (stripped.startswith("**Author:**") or stripped.startswith("**Report:**")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, stripped, size=BODY_PT - 1)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, stripped, size=BODY_PT)
        i += 1

    doc.save(OUT)
    try:
        print(f"wrote {OUT.relative_to(ROOT)}")
    except ValueError:
        print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
